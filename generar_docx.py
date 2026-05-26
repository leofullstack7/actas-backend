from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re

ASSETS_DIR = Path(__file__).resolve().parent / "assets"
LOGO_PATH = ASSETS_DIR / "logo_concejo.png"

MESES = {
    "enero": "Enero",
    "febrero": "Febrero",
    "marzo": "Marzo",
    "abril": "Abril",
    "mayo": "Mayo",
    "junio": "Junio",
    "julio": "Julio",
    "agosto": "Agosto",
    "septiembre": "Septiembre",
    "octubre": "Octubre",
    "noviembre": "Noviembre",
    "diciembre": "Diciembre",
}


def extraer_metadatos_acta(texto: str) -> dict:
    """Extrae tipo de sesión, número de acta y fecha del texto del acta."""
    fragmento = " ".join(texto.split())[:6000]

    sesion = re.search(
        r"sesi[oó]n\s+(ordinaria|extraordinaria|plenaria|especial|de\s+comisi[oó]n)",
        fragmento,
        re.IGNORECASE,
    )
    tipo_sesion = f"Sesión {sesion.group(1).capitalize()}" if sesion else "Sesión ordinaria"

    acta = re.search(
        r"[Aa]cta\s+N[°º.]?\s*0*(\d+)",
        fragmento,
    )
    acta_linea = f"Acta N°. {acta.group(1).zfill(3)}" if acta else "Acta N°. —"

    fecha = re.search(
        r"del\s+(?:[\wáéíóúñ\s\-]+\(\s*)?(\d{1,2})(?:\s*\))?\s+de\s+"
        r"(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+de\s+"
        r"(?:[\wáéíóúñ\s\-]+\(\s*)?(\d{4})",
        fragmento,
        re.IGNORECASE,
    )
    if fecha:
        dia, mes, anio = fecha.group(1), fecha.group(2).lower(), fecha.group(3)
        fecha_linea = f"{MESES.get(mes, mes.capitalize())} {dia} de {anio}"
    else:
        fecha_linea = ""

    return {
        "tipo_sesion": tipo_sesion,
        "acta": acta_linea,
        "fecha": fecha_linea,
    }


def _sin_bordes_tabla(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def agregar_encabezado(doc: Document, metadatos: dict):
    table = doc.add_table(rows=1, cols=2)
    _sin_bordes_tabla(table)
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(12)

    cell_logo, cell_titulo = table.rows[0].cells[0], table.rows[0].cells[1]

    if LOGO_PATH.is_file():
        p_logo = cell_logo.paragraphs[0]
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=Cm(2.6))
    else:
        p_logo = cell_logo.paragraphs[0]
        p_logo.add_run("[Logo: assets/logo_concejo.png]").font.size = Pt(9)

    p1 = cell_titulo.paragraphs[0]
    r1 = p1.add_run("Concejo de")
    r1.font.name = "Arial"
    r1.font.size = Pt(20)
    r1.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    p2 = cell_titulo.add_paragraph()
    r2 = p2.add_run("Manizales")
    r2.font.name = "Arial"
    r2.font.size = Pt(24)
    r2.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    doc.add_paragraph()

    for linea in (metadatos["tipo_sesion"], metadatos["acta"], metadatos["fecha"]):
        if not linea:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(linea)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(16)


def texto_a_docx(texto: str, output_path: str):
    doc = Document()
    metadatos = extraer_metadatos_acta(texto)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = Pt(14)

    agregar_encabezado(doc, metadatos)

    lineas = texto.strip().split("\n")

    for linea in lineas:
        linea = linea.strip()
        if not linea:
            doc.add_paragraph()
            continue

        viñeta = re.match(r"^[-*•❖✓✔►]\s+(.+)", linea)
        numerado = re.match(r"^\d+\.\s+(.+)", linea)
        titulo_punto = re.match(r"^(\d+)\.\t(.+)", linea)

        if titulo_punto:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(f"{titulo_punto.group(1)}.\t{titulo_punto.group(2)}")
            run.bold = True
            run.font.size = Pt(11)

        elif viñeta:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(4)
            aplicar_formato_inline(p, viñeta.group(1))

        elif numerado:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(4)
            aplicar_formato_inline(p, numerado.group(1))

        else:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            aplicar_formato_inline(p, linea)

    doc.save(output_path)


def aplicar_formato_inline(p, texto):
    partes = re.split(r"\*\*(.+?)\*\*", texto)
    for i, parte in enumerate(partes):
        if not parte:
            continue
        run = p.add_run(parte)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
